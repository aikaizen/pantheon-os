from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)

# ── Palette ───────────────────────────────────────────────────────────────────
BLACK      = RGBColor(0x12, 0x12, 0x12)
CHARCOAL   = RGBColor(0x2C, 0x2C, 0x2C)
ACCENT     = RGBColor(0x1A, 0x1A, 0x2E)
MID        = RGBColor(0x55, 0x55, 0x75)
RULE_COLOR = RGBColor(0xCC, 0xCC, 0xD6)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_HEAD = RGBColor(0x1A, 0x1A, 0x2E)
TABLE_ALT  = RGBColor(0xF6, 0xF6, 0xFA)
CALLOUT_BG = RGBColor(0xEC, 0xEC, 0xF4)

BODY_FONT = "Georgia"
HEAD_FONT = "Georgia"

# ── Helpers ───────────────────────────────────────────────────────────────────
def rgb_hex(c):
    return '%02X%02X%02X' % (c[0], c[1], c[2])

def set_font(run, name, size, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def set_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_rule(doc, color=RULE_COLOR):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '4')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), rgb_hex(color))
    pBdr.append(bot)
    pPr.append(pBdr)
    set_spacing(p, before=2, after=2)
    return p

def shade_cell(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  rgb_hex(fill))
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, TABLE_HEAD)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        set_font(run, HEAD_FONT, 9, bold=True, color=WHITE)
        set_spacing(cell.paragraphs[0], before=4, after=4)
    # data
    for ri, row in enumerate(rows):
        fill = TABLE_ALT if ri % 2 == 0 else WHITE
        for ci, val in enumerate(row):
            cell = table.rows[ri+1].cells[ci]
            shade_cell(cell, fill)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, BODY_FONT, 9, color=CHARCOAL)
            set_spacing(cell.paragraphs[0], before=3, after=3)
    if col_widths:
        for row in table.rows:
            for ci, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci])
    doc.add_paragraph()
    return table

def callout(doc, text):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    shade_cell(cell, CALLOUT_BG)
    cell.paragraphs[0].clear()
    run  = cell.paragraphs[0].add_run(text)
    set_font(run, BODY_FONT, 10, italic=True, color=ACCENT)
    set_spacing(cell.paragraphs[0], before=6, after=6)
    cell.paragraphs[0].paragraph_format.left_indent = Pt(8)
    doc.add_paragraph()

def h1(doc, title, number=None):
    p = doc.add_paragraph()
    if number:
        r = p.add_run(f"{number}  ")
        set_font(r, HEAD_FONT, 13, bold=True, color=RULE_COLOR)
    r = p.add_run(title.upper())
    set_font(r, HEAD_FONT, 13, bold=True, color=ACCENT)
    set_spacing(p, before=18, after=6)
    add_rule(doc)

def h2(doc, title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    set_font(r, HEAD_FONT, 10.5, bold=True, color=CHARCOAL)
    set_spacing(p, before=10, after=3)

def body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, BODY_FONT, 10, color=CHARCOAL)
    set_spacing(p, before=2, after=4, line=14)

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    set_font(r, BODY_FONT, 10, color=CHARCOAL)
    set_spacing(p, before=1, after=2, line=13)
    p.paragraph_format.left_indent       = Pt(16 + level * 16)
    p.paragraph_format.first_line_indent = Pt(-10)

# ─────────────────────────────────────────────────────────────────────────────
# COVER
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
r = p.add_run("[COMPANY NAME]")
set_font(r, HEAD_FONT, 28, bold=True, color=ACCENT)
set_spacing(p, before=0, after=4)

p = doc.add_paragraph()
r = p.add_run("Enterprise Agent Systems  ·  Built for Targeted Workloads")
set_font(r, BODY_FONT, 11, italic=True, color=MID)
set_spacing(p, before=0, after=6)

add_rule(doc)

p = doc.add_paragraph()
r = p.add_run("Executive Summary  ·  March 2026  ·  Confidential")
set_font(r, BODY_FONT, 8.5, italic=True, color=MID)
set_spacing(p, before=2, after=20)

# ─────────────────────────────────────────────────────────────────────────────
# 01 WHAT WE ARE
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "What We Are", "01")

body(doc, "We build and manage enterprise agent systems — and we are specifically good at configuring those systems for maximal value in targeted, high-impact workloads.")

body(doc, "This is a meaningful distinction. Most enterprise AI deployments give employees access to general-purpose AI tools and measure adoption. We do something different: we identify the specific workflows where an optimally configured agent system produces a step-change in output, then we engineer and deploy that system. The result is a dramatic improvement in productivity per employee, achieved with minimal token consumption — because the system is designed precisely for the work, not built broadly and hoped for.")

callout(doc, "The resource that matters in agent systems is not compute in the abstract — it is tokens. We design for the best possible output per token, which means the right mix of agents, the right memory architecture, and the right orchestration for each workload. This is not a general AI platform. It is precision engineering for enterprise operational leverage.")

body(doc, "Our forward deployment teams configure these systems inside the client's environment, working close to real workflows until the system is embedded, performing, and trusted. That proximity is how adoption actually happens in large organizations — and it is what separates a working deployment from a failed pilot.")

# ─────────────────────────────────────────────────────────────────────────────
# 02 THE CORE INSIGHT
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "The Core Insight", "02")

body(doc, "The dominant model in enterprise AI — give employees AI tools, watch adoption metrics — produces modest, diffuse, hard-to-measure gains. It also produces significant and often invisible cost, because general-purpose agents are expensive to run at scale and most of that spend produces marginal value.")

body(doc, "The alternative is workload-targeted deployment: identify where the operational leverage is genuinely high, configure an agent system purpose-built for that workload, and measure productivity per employee in that workflow directly. This approach produces:")

bullet(doc, "Larger, more defensible productivity gains — because the system is optimized for the specific task")
bullet(doc, "Lower token consumption — because the agent doesn't need to be capable of everything, only excellent at one thing")
bullet(doc, "Faster proof of value — because outcomes are measurable and localized, not diffuse")
bullet(doc, "Higher trust — because the system works predictably within a defined scope")

body(doc, "Before a targeted agent system can work reliably in an enterprise context, three infrastructure layers must be in place:")

add_table(doc,
    headers=["Layer", "Question it answers", "What we build"],
    rows=[
        ["Memory",   "What does the organization know, and can the agent access it safely?",        "Institutional Memory Engine — structured, permission-aware, cited"],
        ["Workflow",  "How does the work actually move, and where does it stall?",                  "Long-Horizon Workflow Agent — state-tracked, multi-week, escalation-aware"],
        ["Control",   "Who governs what the agents can do, and what did they do?",                  "Agent Runtime Control Plane — approval gates, audit trail, cost visibility"],
    ],
    col_widths=[1.2, 3.0, 2.4]
)

body(doc, "These three layers are not three separate products sold independently. They are the substrate — the foundation that makes a targeted agent system reliable, governable, and cumulative. We configure and deploy them together, calibrated to the specific workload and the specific organization.")

# ─────────────────────────────────────────────────────────────────────────────
# 03 WHAT WE BUILD
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "What We Build & Deploy", "03")

h2(doc, "Workload-Targeted Agent Systems")
body(doc, "The primary deliverable is a configured, deployed, and managed agent system optimized for one or more specific enterprise workloads. Not a platform. Not a general copilot. A precision system for a defined operational problem — onboarding at scale, compliance coordination, delivery program management, shared services operations — where we can engineer a measurable productivity gain.")

h2(doc, "The Memory Layer")
body(doc, "Every agent system we deploy is backed by a structured memory architecture. This captures what the organization knows — decisions, workflows, client context, institutional patterns — in a form that agents can query reliably and that humans can govern. Without structured memory, agents are stateless and expensive. With it, they become cumulative, cheaper per query, and significantly more accurate.")

h2(doc, "The Orchestration Layer")
body(doc, "We configure the right mix of agents for each workload. This is one of the most consequential engineering decisions in any deployment: which tasks are handled autonomously, which require human-in-the-loop approval, which agents hand off to other agents, and at what token cost. We optimize this mix for the specific workload, not for generality.")

h2(doc, "The Control Layer")
body(doc, "We deploy a control plane that gives the organization full visibility and governance over every agent running in its environment — what is active, what it touched, what decisions it made, what it cost. This is what makes enterprise adoption trustworthy rather than experimental.")

h2(doc, "Forward Deployment")
body(doc, "Our teams configure and embed these systems inside the client's environment. This is the adoption model — not just the delivery model. We stay close until the system is performing, the team trusts it, and the productivity gain is measurable. Then we step back into a managed service or subscription posture.")

# ─────────────────────────────────────────────────────────────────────────────
# 04 WHAT WE ARE NOT
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "What We Are Not", "04")

add_table(doc,
    headers=["We are", "We are not"],
    rows=[
        ["Enterprise agent system builders — precise, workload-targeted",        "A platform company selling generic AI tools to employees"],
        ["Optimizers of productivity per employee and output per token",          "An AI vendor measuring seats, adoption rates, or prompt volume"],
        ["Forward-deployed configurators of memory, orchestration, and control", "A consultancy selling hours or a systems integrator adding headcount"],
        ["Sold on measurable productivity gains in specific workflows",           "Sold on AI novelty, capability demos, or technology access"],
        ["Engineers of the right agent mix for each workload",                   "Builders of general-purpose copilots or broad employee tools"],
    ],
    col_widths=[3.3, 3.3]
)

# ─────────────────────────────────────────────────────────────────────────────
# 05 INFRASTRUCTURE PARTNERSHIP
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "Infrastructure Partnership — On-Prem & Multicloud", "05")

body(doc, "A strategic infrastructure partnership is currently in formation with an undisclosed enterprise technology partner. This will enable on-premises and multicloud deployment of the full agent system stack — memory, orchestration, and control plane — within the client's own environment.")

body(doc, "This is commercially significant for several reasons:")

bullet(doc, "Data residency requirements in Saudi Arabia and other regulated environments frequently make cloud-only deployments a blocker at the procurement stage. On-prem removes this objection before it is raised.")
bullet(doc, "Government-adjacent and regulated-sector accounts in the Gulf commonly require sovereign or in-country deployment. This partnership positions us to pursue those accounts from the first conversation.")
bullet(doc, "Token efficiency is further improved in on-prem configurations where the organization controls the runtime environment and can optimize for workload-specific throughput.")
bullet(doc, "The partnership may open co-sell and referral pipeline through an established enterprise infrastructure network — a significant distribution advantage.")

body(doc, "The product architecture is being designed to support on-prem deployment from the outset. This is not a later-stage retrofit.")

# ─────────────────────────────────────────────────────────────────────────────
# 06 GO-TO-MARKET
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "Go-to-Market Model", "06")

h2(doc, "The Delivery Sequence")

add_table(doc,
    headers=["Phase", "Description", "Output"],
    rows=[
        ["01  Diagnostic Sprint",               "Paid, fixed-scope. Map the target workload, quantify the productivity gap, model the agent system configuration and expected ROI.",                         "Trust + scoped proposal"],
        ["02  Fixed-Scope Pilot",               "One workload, one configured agent system, success criteria agreed upfront. Measurable productivity per employee outcome.",                                 "Proof of value"],
        ["03  Forward-Deployed Implementation", "Embedded team configures and operationalizes memory, orchestration, and control inside the client environment. Stays until system performs.",               "Adopted, trusted system"],
        ["04  Managed Service + Subscription",  "Ongoing system management, optimization, and governance. Subscription tied to workload complexity, not seat count.",                                       "Recurring revenue + expansion"],
    ],
    col_widths=[2.0, 3.5, 1.1]
)

h2(doc, "Ideal Customer Profile")
body(doc, "Best-fit buyers are large enterprises with high-volume, high-stakes operational workflows — onboarding at scale, delivery program coordination, compliance management, shared services operations — where a 20–40% productivity gain per employee in a specific function produces material business impact. They understand the cost of coordination failure and are willing to invest in a system that solves it precisely.")

body(doc, "We do not pursue broad employee-tool deployments, general AI platform evaluations, or buyers without a specific operational problem and mandate to solve it.")

# ─────────────────────────────────────────────────────────────────────────────
# 07 PRICING
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "Pricing Direction", "07")

h2(doc, "Principles")
bullet(doc, "Price on productivity outcomes, not on hours, seats, or tokens consumed.")
bullet(doc, "Every engagement tier must cover forward deployment cost. No subsidizing of discovery.")
bullet(doc, "Subscription pricing reflects the complexity of the deployed workload and agent system — not a simple per-seat model.")

h2(doc, "Indicative Pricing")
add_table(doc,
    headers=["Offer", "Commercial Model", "Indicative Range"],
    rows=[
        ["Diagnostic Sprint",               "Fixed fee",                              "$15,000 – $30,000"],
        ["Pilot Implementation",            "Fixed fee, success-scoped",               "$40,000 – $100,000"],
        ["Forward-Deployed Implementation", "Time-and-scope, per engagement",          "$75,000 – $250,000"],
        ["Annual Software Subscription",    "Per workload + system complexity tier",   "$120,000 – $600,000+"],
        ["Managed Service Retainer",        "Monthly recurring",                       "$20,000 – $60,000 / month"],
    ],
    col_widths=[2.4, 2.4, 1.8]
)

h2(doc, "Subscription Metric")
body(doc, "Annual subscription pricing is tied to the complexity of the deployed agent system — number of active workloads under management, agent configurations, memory workspace size, and operator seats — rather than a flat per-user model. This aligns our commercial incentives with the operational depth we create, not with headcount.")

h2(doc, "Revenue Mix Target")
add_table(doc,
    headers=["Phase", "Timeline", "Services", "Recurring"],
    rows=[
        ["Phase 1 — Current",    "Now",           "80%", "20%"],
        ["Phase 2 — Growth",     "12–18 months",  "55%", "45%"],
        ["Phase 3 — Scale",      "24–36 months",  "30%", "70%"],
    ],
    col_widths=[1.9, 1.5, 1.4, 1.8]
)

body(doc, "Services remain essential — they are how we discover the right workload, configure the system correctly, and make adoption happen in complex enterprise environments. The goal is to make services feed the product roadmap and shift toward higher-margin advisory work as the recurring base grows.")

# ─────────────────────────────────────────────────────────────────────────────
# 08 SAUDI ARABIA & MIDDLE EAST
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "Saudi Arabia & the Middle East", "08")

callout(doc, "Saudi Arabia is the primary strategic market — not a later-stage expansion. The combination of large-scale operational transformation mandates, forward deployment as the natural enterprise sales motion, and on-prem deployment capability makes it the strongest initial fit for this model.")

h2(doc, "Why Saudi First")
bullet(doc, "Vision 2030 has produced the largest enterprise transformation budget in the region, with explicit mandates in logistics, real estate, healthcare, and government-adjacent operations — each a strong fit for targeted agent system deployment.")
bullet(doc, "Enterprise sales in Saudi Arabia is relationship-led, high-trust, and favor embedded delivery over remote software. This maps precisely to the forward deployment model.")
bullet(doc, "On-prem and sovereign deployment options — enabled by the infrastructure partnership — remove the single largest technical objection in regulated sectors and government-adjacent accounts.")
bullet(doc, "Deal sizes and contract terms in the Saudi enterprise market support the economics of forward deployment. A single well-structured account justifies the investment.")
bullet(doc, "The agent efficiency narrative — measurable productivity per employee with controlled token cost — is directly legible to Saudi CFOs and transformation officers managing large implementation budgets.")

h2(doc, "Regional Approach")
body(doc, "Saudi Arabia is the primary strategic target. The UAE functions as the regional commercial infrastructure — faster pilots, English-led sales, international partner relationships, and the base from which Gulf references are built. UAE speed builds the credibility that Saudi relationships require. Both are active in parallel; Saudi is not a reward after UAE success.")

h2(doc, "Priority Sectors")
add_table(doc,
    headers=["Sector", "Fit with targeted agent deployment"],
    rows=[
        ["Logistics & supply chain",          "High-volume, multi-team coordination; measurable throughput gains; strong ROI model"],
        ["Real estate & construction delivery","Multi-week delivery programs with complex stakeholder chains; milestone-dependent workflows"],
        ["Healthcare operations",             "Non-clinical coordination at scale; strong data residency requirements; on-prem a prerequisite"],
        ["Government-adjacent programs",      "Large transformation mandates; sovereign deployment required; high willingness to invest"],
        ["Large family business groups",      "Multi-line-of-business operational complexity; high ACV potential; relationship-led buying"],
        ["Enterprise shared services",        "High-volume coordination overhead; strong memory and orchestration fit; measurable time savings"],
    ],
    col_widths=[2.3, 4.3]
)

h2(doc, "Operational Requirements")
bullet(doc, "Arabic executive materials — one-pager, proposal template, and executive summary before major outreach.")
bullet(doc, "On-prem and data residency options confirmed before regulated-sector or government-adjacent procurement conversations.")
bullet(doc, "Local legal entity or credible in-country partner before large public sector or regulated-sector work begins.")
bullet(doc, "After-sales support model that matches regional expectations. Face-to-face presence still carries disproportionate weight at the senior level.")

# ─────────────────────────────────────────────────────────────────────────────
# 09 WHAT WE DO NEXT
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "What We Do Next", "09")

h2(doc, "Immediate — Next 30 Days")
bullet(doc, "Finalize company name and core positioning statement as an enterprise agent system builder, not a consultancy.")
bullet(doc, "Package the Diagnostic Sprint as a replicable, fixed-fee offer with a standard workload-mapping output template.")
bullet(doc, "Define the first productized system: Institutional Memory Engine — scope, configuration playbook, pricing.")
bullet(doc, "Identify 3–5 priority Saudi target accounts or sectors and begin relationship mapping.")
bullet(doc, "Finalize terms and architecture implications of the infrastructure partnership.")

h2(doc, "Near-Term — 30–60 Days")
bullet(doc, "Build the Memory Engine as an internal tool first, then as a client-deployable module.")
bullet(doc, "Create the standard forward-deployed implementation playbook for the first workload type.")
bullet(doc, "Develop sector-specific materials for two priority Saudi verticals.")
bullet(doc, "Produce Arabic executive summary and one-pager.")
bullet(doc, "Run one direct outbound motion to Saudi enterprise targets.")

h2(doc, "Medium-Term — 60–90 Days")
bullet(doc, "Convert the first repeated workload configuration into a replicable product module.")
bullet(doc, "Publish one case study with specific productivity-per-employee and token-efficiency metrics.")
bullet(doc, "Identify Saudi partner candidates — systems integrator or local operator with enterprise procurement access.")
bullet(doc, "Prepare data residency and compliance documentation for regional enterprise deals.")
bullet(doc, "Begin Long-Horizon Workflow Agent buildout for multi-week coordination workloads.")

# ─────────────────────────────────────────────────────────────────────────────
# 10 KEY RISKS
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "Key Risks", "10")

add_table(doc,
    headers=["Risk", "Mitigation"],
    rows=[
        ["Drifting back toward general consultancy",          "Productization review after every engagement; maintain workload-specific product roadmap independent of client requests"],
        ["Productizing before repeatability is proven",       "Rule: only standardize configurations that have been deployed across three or more distinct engagements"],
        ["Token cost narrative misunderstood by buyers",      "Lead with productivity-per-employee outcomes; introduce token efficiency as the underlying mechanism, not the headline"],
        ["Forward deployment economics becoming unsustainable","Diagnostic and pilot phases priced to fully cover cost; no subsidizing of configuration discovery"],
        ["Infrastructure partnership delay",                  "Architecture designed for on-prem from day one; partnership accelerates deal flow but does not gate the product"],
        ["Saudi enterprise sales cycles longer than modeled", "UAE used for faster pilots and reference-building; Saudi pursued in parallel with appropriate relationship investment"],
    ],
    col_widths=[2.8, 3.8]
)

# ─────────────────────────────────────────────────────────────────────────────
# CLOSING
# ─────────────────────────────────────────────────────────────────────────────
h1(doc, "Summary", "")

callout(doc, "We build and manage enterprise agent systems configured for maximal value in targeted workloads — not general AI tools for employees, but precision systems that produce a step-change in productivity per employee with minimal token consumption. Three infrastructure layers — memory, orchestration, control — deployed by forward teams, measured on operational outcomes, and anchored in Saudi Arabia as the primary strategic market. On-prem deployment capability removes the region's largest technical objection from day one.")

doc.add_paragraph()
add_rule(doc)
p = doc.add_paragraph()
r = p.add_run("[Company Name]  ·  Confidential  ·  March 2026")
set_font(r, BODY_FONT, 8, italic=True, color=MID)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_spacing(p, before=6, after=0)

# ─────────────────────────────────────────────────────────────────────────────
out = "/Users/adilislam/Desktop/Hermes/portal/consultancy-to-product-strategy/Executive-Summary-March-2026.docx"
doc.save(out)
print(f"Saved: {out}")
